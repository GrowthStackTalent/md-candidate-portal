from __future__ import annotations

import copy
import csv
import os
import re
import smtplib
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from functools import wraps
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from urllib.parse import quote, unquote
from flask import (Flask, render_template, request, jsonify,
                   send_from_directory, session, redirect, url_for)
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-change-in-production")
app.jinja_env.filters["urlquote"] = lambda s: quote(str(s), safe="")
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024

limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=[],
    storage_uri="memory://",
)

BASE_DIR = Path(__file__).parent
DATA_DIR = Path(os.environ.get("DATA_DIR", BASE_DIR))
CSV_PATH = DATA_DIR / "candidates.csv"
JOBS_PATH = DATA_DIR / "jobs.csv"
ASSIGNMENTS_PATH = DATA_DIR / "assignments.csv"
CLIENTS_PATH = DATA_DIR / "clients.csv"
TAGS_PATH = DATA_DIR / "tags.csv"
CANDIDATE_TAGS_PATH = DATA_DIR / "candidate_tags.csv"
OUTPUT_DIR = DATA_DIR / "output"
UPLOADS_DIR = DATA_DIR / "uploads"
TERMS_DIR = DATA_DIR / "terms"
OFFER_LETTERS_DIR = DATA_DIR / "offer_letters"
OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
UPLOADS_DIR.mkdir(exist_ok=True, parents=True)
TERMS_DIR.mkdir(exist_ok=True, parents=True)
OFFER_LETTERS_DIR.mkdir(exist_ok=True, parents=True)

ALLOWED_CV_EXTENSIONS = {".pdf", ".doc", ".docx"}

CSV_HEADERS = [
    "Submission Date", "Full Name", "Email", "Phone", "Location", "LinkedIn URL",
    "Current Job Title", "Current Company", "Currency",
    "Current Base Salary", "On Target Earnings",
    "Benefits", "Open to Relocate", "Nearest Airport",
    "Languages", "Education", "Notice Period", "CV / Resume", "CV Filename",
    "Word Doc", "Data Consent Given", "Job Applied For",
]

JOBS_HEADERS = ["Job ID", "Date Added", "Company Name", "Job Title", "Location",
                "Salary Band", "Currency", "Contact Person", "Assigned To", "Status",
                "Description", "Advert Text", "Show on Website",
                "Expected Start Date", "Expected Invoice Date"]
ASSIGNMENTS_HEADERS = ["Job ID", "Candidate Date", "Candidate Name", "Stage", "Offer Letter Filename"]

PIPELINE_STAGES = [
    "First Interview", "Second Interview", "Third Interview", "Fourth Interview",
    "Offer", "Accepted", "Declined", "Removed from Process",
]

STAGE_COLOURS = {
    "First Interview":      "blue",
    "Second Interview":     "blue",
    "Third Interview":      "blue",
    "Fourth Interview":     "blue",
    "Offer":                "orange",
    "Accepted":             "green",
    "Declined":             "red",
    "Removed from Process": "grey",
}
CLIENTS_HEADERS = ["Client ID", "Date Added", "Client Name", "Contact Name", "Fee Percentage", "Terms Filename"]
TAGS_HEADERS = ["Tag ID", "Tag Name", "Colour"]
CANDIDATE_TAGS_HEADERS = ["Candidate Date", "Tag Name"]

TAG_COLOURS = ["blue", "green", "orange", "purple", "red", "teal", "yellow", "indigo"]

CURRENCY_SYMBOLS = {
    "GBP": "£", "USD": "$", "EUR": "€", "CAD": "CA$",
    "AUD": "A$", "CHF": "CHF ", "JPY": "¥", "SGD": "S$",
    "HKD": "HK$", "AED": "AED ", "INR": "₹",
}


# ── CSV helpers ────────────────────────────────────────────────────────────────

def ensure_csv():
    if not CSV_PATH.exists():
        with open(CSV_PATH, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(CSV_HEADERS)


def append_to_csv(row: dict):
    ensure_csv()
    with open(CSV_PATH, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=CSV_HEADERS, extrasaction="ignore")
        writer.writerow(row)


def ensure_jobs_csv():
    if not JOBS_PATH.exists():
        with open(JOBS_PATH, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(JOBS_HEADERS)

def ensure_assignments_csv():
    if not ASSIGNMENTS_PATH.exists():
        with open(ASSIGNMENTS_PATH, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(ASSIGNMENTS_HEADERS)

def ensure_clients_csv():
    if not CLIENTS_PATH.exists():
        with open(CLIENTS_PATH, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(CLIENTS_HEADERS)

def read_all_clients() -> list[dict]:
    ensure_clients_csv()
    with open(CLIENTS_PATH, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def ensure_tags_csv():
    if not TAGS_PATH.exists():
        with open(TAGS_PATH, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(TAGS_HEADERS)

def ensure_candidate_tags_csv():
    if not CANDIDATE_TAGS_PATH.exists():
        with open(CANDIDATE_TAGS_PATH, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(CANDIDATE_TAGS_HEADERS)

def read_all_tags() -> list[dict]:
    ensure_tags_csv()
    with open(TAGS_PATH, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def read_candidate_tags() -> list[dict]:
    ensure_candidate_tags_csv()
    with open(CANDIDATE_TAGS_PATH, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def read_all_jobs() -> list[dict]:
    ensure_jobs_csv()
    with open(JOBS_PATH, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def read_assignments() -> list[dict]:
    ensure_assignments_csv()
    with open(ASSIGNMENTS_PATH, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))

def read_all_candidates() -> list[dict]:
    if not CSV_PATH.exists():
        return []
    with open(CSV_PATH, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    rows.reverse()
    return rows


# ── Utilities ──────────────────────────────────────────────────────────────────

def _build_job_location(city: str, country: str, remote: bool) -> str:
    if remote:
        return f"Remote — {city}, {country}" if city else f"Remote — {country}" if country else "Remote"
    return f"{city}, {country}" if city and country else city or country or "—"


def safe_filename(name: str) -> str:
    return re.sub(r"[^\w\-]", "_", name)


def save_cv(file, candidate_name: str) -> tuple[str, str]:
    original_name = file.filename or ""
    ext = Path(original_name).suffix.lower()
    if ext not in ALLOWED_CV_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Please upload a PDF or Word document.")
    saved_name = f"{safe_filename(candidate_name)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_CV{ext}"
    file.save(UPLOADS_DIR / saved_name)
    return original_name, saved_name


# ── Email notification ─────────────────────────────────────────────────────────

def send_notification_email(data: dict):
    smtp_host = os.environ.get("SMTP_HOST", "")
    if not smtp_host:
        return
    try:
        smtp_port = int(os.environ.get("SMTP_PORT", "587"))
        smtp_user = os.environ.get("SMTP_USER", "")
        smtp_password = os.environ.get("SMTP_PASSWORD", "")
        recruiter_emails_raw = os.environ.get("RECRUITER_EMAIL", smtp_user)
        recruiter_emails = [e.strip() for e in recruiter_emails_raw.split(",") if e.strip()]

        msg = MIMEMultipart()
        msg["Subject"] = f"New Application: {data['Full Name']} — {data['Current Job Title']}"
        msg["From"] = smtp_user
        msg["To"] = ", ".join(recruiter_emails)

        cv_line = f"Yes — {data['CV / Resume']}" if data.get("CV / Resume") else "Not uploaded"
        body = "\n".join([
            "New candidate application received via Mason Dynamics LTD portal.",
            "",
            f"Name:           {data['Full Name']}",
            f"Location:       {data['Location']}",
            f"LinkedIn:       {data['LinkedIn URL']}",
            f"Job Title:      {data['Current Job Title']}",
            f"Company:        {data['Current Company']}",
            f"Base Salary:    {data['Current Base Salary']}",
            f"OTE:            {data['On Target Earnings']}",
            f"Benefits:       {data['Benefits']}",
            f"Notice Period:  {data['Notice Period']}",
            f"Relocate:       {data['Open to Relocate']}",
            f"Languages:      {data['Languages']}",
            f"CV Uploaded:    {cv_line}",
            f"Submitted:      {data['Submission Date']}",
            "",
            "Log in to the admin portal to view all submissions and download documents.",
        ])
        msg.attach(MIMEText(body, "plain"))

        # Attach Word front sheet
        doc_path = OUTPUT_DIR / data.get("Word Doc", "")
        if data.get("Word Doc") and doc_path.exists():
            with open(doc_path, "rb") as f:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition",
                            f'attachment; filename="{data["Full Name"]} - Front Sheet.docx"')
            msg.attach(part)

        # Attach CV if uploaded
        cv_path = UPLOADS_DIR / data.get("CV Filename", "")
        if data.get("CV Filename") and cv_path.exists():
            with open(cv_path, "rb") as f:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition",
                            f'attachment; filename="{data["CV / Resume"]}"')
            msg.attach(part)

        raw = msg.as_string()
        if smtp_port == 465:
            with smtplib.SMTP_SSL(smtp_host, smtp_port) as server:
                server.login(smtp_user, smtp_password)
                for recipient in recruiter_emails:
                    msg.replace_header("To", recipient)
                    server.sendmail(smtp_user, recipient, msg.as_string())
        else:
            with smtplib.SMTP(smtp_host, smtp_port) as server:
                server.ehlo()
                server.starttls()
                server.login(smtp_user, smtp_password)
                for recipient in recruiter_emails:
                    msg.replace_header("To", recipient)
                    server.sendmail(smtp_user, recipient, msg.as_string())
    except Exception as e:
        app.logger.error("Email notification failed: %s", e)


# ── Word document ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


FONT = "Helvetica"


def _append_docx(bio_doc, cv_path: Path):
    """Append all body content from a .docx CV onto the bio document after a page break."""
    cv_doc = Document(str(cv_path))
    bio_doc.add_page_break()
    body = bio_doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for elem in cv_doc.element.body:
        if elem.tag == qn("w:sectPr"):
            continue
        if sect_pr is not None:
            sect_pr.addprevious(copy.deepcopy(elem))
        else:
            body.append(copy.deepcopy(elem))


def generate_word_doc(data: dict, cv_path: Path | None = None) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Set default font for the document
    doc.styles["Normal"].font.name = FONT

    logo_path = BASE_DIR / "static" / "logo.png"
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = "Table Grid"

    logo_cell = header_table.cell(0, 0)
    logo_cell.width = Inches(3.2)
    set_cell_bg(logo_cell, "FFFFFF")
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_before = Pt(6)
    logo_para.paragraph_format.space_after = Pt(6)
    if logo_path.exists():
        logo_para.add_run().add_picture(str(logo_path), width=Inches(2.8))

    title_cell = header_table.cell(0, 1)
    title_cell.width = Inches(3.4)
    set_cell_bg(title_cell, "1B3A5C")
    hdr_para = title_cell.paragraphs[0]
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr_para.add_run("CANDIDATE BIO")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr_para.paragraph_format.space_before = Pt(18)
    hdr_para.paragraph_format.space_after = Pt(18)

    doc.add_paragraph()

    fields = [
        ("Full Name",           data["Full Name"]),
        ("Email",               data["Email"]),
        ("Phone",               data["Phone"]),
        ("Location",            data["Location"]),
        ("LinkedIn",            data["LinkedIn URL"]),
        ("Current Job Title",   data["Current Job Title"]),
        ("Current Company",     data["Current Company"]),
        ("Current Base Salary",         data["Current Base Salary"]),
        ("Current On Target Earnings",  data["On Target Earnings"]),
        ("Current Benefits",            data["Benefits"]),
        ("Open to Relocate",    data["Open to Relocate"]),
        ("Nearest Airport",     data["Nearest Airport"]),
        ("Languages",           data["Languages"]),
        ("Education",           data["Education"]),
        ("Notice Period",       data["Notice Period"]),
    ]

    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Inches(2.1), Inches(4.5)]

    for i, (label, value) in enumerate(fields):
        row = tbl.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        label_cell = row.cells[0]
        set_cell_bg(label_cell, "E8EDF3")
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = FONT
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

        value_cell = row.cells[1]
        set_cell_bg(value_cell, "FFFFFF")
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after = Pt(4)
        vr = vp.add_run(str(value) if value else "—")
        vr.font.name = FONT
        vr.font.size = Pt(10)

    # ── Interview notes section ────────────────────────────────────────────────
    doc.add_paragraph()

    notes_heading = doc.add_paragraph()
    notes_heading.paragraph_format.space_before = Pt(10)
    notes_heading.paragraph_format.space_after = Pt(6)
    nh_run = notes_heading.add_run("CANDIDATE INTERVIEW NOTES")
    nh_run.bold = True
    nh_run.font.name = FONT
    nh_run.font.size = Pt(10)
    nh_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # ── Page footer ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.footer_distance = Cm(1)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr1 = fp.add_run(
            "All candidate submissions are deemed acceptance to our standard terms & conditions "
            "unless otherwise stated in writing\n"
            "GrowthStack Consultants LLC  •  info@masondynamics.com  •  www.masondynamics.com"
        )
        fr1.font.name = FONT
        fr1.font.size = Pt(8)
        fr1.font.color.rgb = RGBColor(0x4A, 0x5A, 0x6A)
        fr1.italic = True

    if cv_path and cv_path.exists() and cv_path.suffix.lower() == ".docx":
        _append_docx(doc, cv_path)

    fname = f"{safe_filename(data['Full Name'])}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(OUTPUT_DIR / fname)
    return fname


# ── Auth decorator ─────────────────────────────────────────────────────────────

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("admin_logged_in"):
            return redirect(url_for("admin_login"))
        return f(*args, **kwargs)
    return decorated


# ── Public routes ──────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("form.html")


@app.route("/submit", methods=["POST"])
def submit():
    f = request.form

    benefits_selected = f.getlist("benefits")
    other_benefit = f.get("benefits_other", "").strip()
    if other_benefit and "Other" in benefits_selected:
        benefits_selected = [b for b in benefits_selected if b != "Other"]
        benefits_selected.append(f"Other: {other_benefit}")
    benefits_str = ", ".join(benefits_selected) if benefits_selected else "None"

    lang_names = f.getlist("lang_name")
    lang_levels = f.getlist("lang_level")
    langs = [f"{n} ({l})" for n, l in zip(lang_names, lang_levels) if n.strip()]
    languages_str = ", ".join(langs) if langs else "—"

    edu_institutions = f.getlist("edu_institution")
    edu_degrees = f.getlist("edu_degree")
    edu_subjects = f.getlist("edu_subject")
    edu_years = f.getlist("edu_year")
    edu_entries = []
    for inst, deg, subj, yr in zip(edu_institutions, edu_degrees, edu_subjects, edu_years):
        if inst.strip():
            parts = [deg.strip(), subj.strip(), inst.strip()]
            if yr.strip():
                parts.append(yr.strip())
            edu_entries.append(" | ".join(p for p in parts if p))
    education_str = "; ".join(edu_entries) if edu_entries else "—"

    currency = f.get("currency", "GBP")
    symbol = CURRENCY_SYMBOLS.get(currency, "")

    def fmt_salary(val):
        try:
            return f"{symbol}{int(float(val)):,}"
        except (ValueError, TypeError):
            return val or "—"

    full_name = f.get("full_name", "").strip()

    job_applied_for = ""
    job_id_applied = f.get("job_id", "").strip()
    if job_id_applied:
        all_jobs = read_all_jobs()
        matched = next((j for j in all_jobs if j["Job ID"] == job_id_applied), None)
        if matched:
            job_applied_for = f"{matched['Job Title']} — {matched['Company Name']}"

    row = {
        "Submission Date":    datetime.now().strftime("%Y-%m-%d %H:%M"),
        "Full Name":          full_name,
        "Email":              f.get("email", "").strip(),
        "Phone":              f.get("phone_code", "").strip() + f.get("phone_number", "").strip(),
        "Location":           f.get("location", "").strip(),
        "LinkedIn URL":       f.get("linkedin", "").strip(),
        "Current Job Title":  f.get("job_title", "").strip(),
        "Current Company":    f.get("company", "").strip(),
        "Currency":           currency,
        "Current Base Salary": fmt_salary(f.get("base_salary")),
        "On Target Earnings":  fmt_salary(f.get("ote")),
        "Benefits":           benefits_str,
        "Open to Relocate":   f.get("relocate", "No"),
        "Nearest Airport":    f.get("airport", "").strip(),
        "Languages":          languages_str,
        "Education":          education_str,
        "Notice Period":      f.get("notice_period", "").strip(),
        "CV / Resume":        "",
        "CV Filename":        "",
        "Word Doc":           "",
        "Data Consent Given": f"Yes – {datetime.now().strftime('%Y-%m-%d %H:%M')} UTC",
        "Job Applied For":    job_applied_for,
    }

    try:
        cv_file = request.files.get("cv_file")
        cv_path_for_merge = None
        if cv_file and cv_file.filename:
            original_name, saved_name = save_cv(cv_file, full_name)
            row["CV / Resume"] = original_name
            row["CV Filename"] = saved_name
            if Path(saved_name).suffix.lower() == ".docx":
                cv_path_for_merge = UPLOADS_DIR / saved_name

        doc_filename = generate_word_doc(row, cv_path_for_merge)
        row["Word Doc"] = doc_filename

        append_to_csv(row)
        send_notification_email(row)

        return jsonify({"status": "ok", "filename": doc_filename})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


@app.route("/jobs")
def public_jobs():
    all_jobs = read_all_jobs()
    open_jobs = [
        j for j in reversed(all_jobs)
        if j.get("Status") == "Open" and j.get("Show on Website", "yes") == "yes"
    ]
    return render_template("jobs.html", jobs=open_jobs)


@app.route("/jobs/<job_id>/apply")
def job_apply(job_id):
    all_jobs = read_all_jobs()
    job = next((j for j in all_jobs if j["Job ID"] == job_id), None)
    if not job or job.get("Status") != "Open":
        return redirect(url_for("public_jobs"))
    currency_symbol = CURRENCY_SYMBOLS.get(job.get("Currency", "GBP") or "GBP", "£")
    return render_template("job_apply.html", job=job, currency_symbol=currency_symbol)


# ── Admin routes ───────────────────────────────────────────────────────────────

@app.route("/admin", methods=["GET", "POST"])
@limiter.limit("5 per 15 minutes", methods=["POST"], error_message="LOCKED")
def admin_login():
    if session.get("admin_logged_in"):
        return redirect(url_for("admin_home"))
    error = None
    if request.method == "POST":
        if request.form.get("password") == os.environ.get("ADMIN_PASSWORD", "admin"):
            session["admin_logged_in"] = True
            return redirect(url_for("admin_home"))
        error = "Incorrect password. Please try again."
    return render_template("admin_login.html", error=error)


@app.route("/admin/home")
@admin_required
def admin_home():
    import re as _re
    candidates  = read_all_candidates()
    jobs        = read_all_jobs()
    assignments = read_assignments()
    clients     = read_all_clients()

    # --- Core counts ---
    # Treat blank/missing Status as "Open" (default when job is first created)
    active_jobs     = sum(1 for j in jobs if j.get("Status", "Open").strip().lower() in ("open", ""))
    filled_jobs     = sum(1 for j in jobs if "successfully filled" in j.get("Status", "").strip().lower())
    total_candidates = len(candidates)
    active_clients  = len(clients)

    # --- Pipeline stage counts (deduplicated per candidate-job pair) ---
    interview_stages = {"First Interview", "Second Interview", "Third Interview", "Fourth Interview"}
    at_interview = sum(1 for a in assignments if a.get("Stage") in interview_stages)
    at_offer     = sum(1 for a in assignments if a.get("Stage") == "Offer")
    accepted     = sum(1 for a in assignments if a.get("Stage") == "Accepted")

    # --- Value of accepted offers (fee = salary * fee%) ---
    # Build lookup maps
    job_map    = {j["Job ID"]: j for j in jobs}
    client_map = {c["Client Name"]: c for c in clients}

    def _parse_salary(band: str) -> float | None:
        """Extract first number from a salary band string e.g. '£60,000 - £80,000'."""
        digits = _re.sub(r"[^\d.]", "", band.split("-")[0].split("–")[0].strip())
        try:
            return float(digits) if digits else None
        except ValueError:
            return None

    accepted_value = 0.0
    for a in assignments:
        if a.get("Stage") != "Accepted":
            continue
        job = job_map.get(a.get("Job ID", ""))
        if not job:
            continue
        salary = _parse_salary(job.get("Salary Band", ""))
        if salary is None:
            continue
        client = client_map.get(job.get("Company Name", ""))
        fee_pct = 0.0
        if client:
            try:
                fee_pct = float(client.get("Fee Percentage", 0))
            except (ValueError, TypeError):
                fee_pct = 0.0
        if fee_pct:
            accepted_value += salary * (fee_pct / 100)

    # --- This month ---
    now = datetime.now()
    month_prefix = now.strftime("%Y-%m")
    candidates_this_month = sum(
        1 for c in candidates if c.get("Submission Date", "").startswith(month_prefix)
    )

    # --- Unplaced candidates (no assignment at all) ---
    assigned_dates = {a["Candidate Date"] for a in assignments}
    unplaced = sum(1 for c in candidates if c.get("Submission Date") not in assigned_dates)

    # --- Recent submissions (last 5) ---
    recent = candidates[:5]

    # --- Greeting ---
    hour = datetime.now().hour
    if hour < 12:
        greeting = "morning"
    elif hour < 17:
        greeting = "afternoon"
    else:
        greeting = "evening"

    return render_template(
        "admin_home.html",
        active_jobs=active_jobs,
        filled_jobs=filled_jobs,
        total_candidates=total_candidates,
        active_clients=active_clients,
        at_interview=at_interview,
        at_offer=at_offer,
        accepted=accepted,
        accepted_value=accepted_value,
        candidates_this_month=candidates_this_month,
        unplaced=unplaced,
        recent=recent,
        greeting=greeting,
    )


@app.errorhandler(429)
def rate_limit_handler(e):
    return render_template("admin_login.html",
                           error="Too many failed attempts. Access locked for 15 minutes."), 429


@app.route("/admin/dashboard")
@admin_required
def admin_dashboard():
    candidates = read_all_candidates()
    jobs = read_all_jobs()
    assignments = read_assignments()
    assigned_jobs = {}
    for a in assignments:
        assigned_jobs.setdefault(a["Candidate Date"], []).append(a)
    tags = read_all_tags()
    candidate_tags = read_candidate_tags()
    tags_by_candidate = {}
    for ct in candidate_tags:
        tags_by_candidate.setdefault(ct["Candidate Date"], []).append(ct["Tag Name"])
    return render_template("admin_dashboard.html", candidates=candidates,
                           jobs=jobs, assigned_jobs=assigned_jobs,
                           tags=tags, tags_by_candidate=tags_by_candidate,
                           stage_colours=STAGE_COLOURS)


@app.route("/admin/logout")
def admin_logout():
    session.pop("admin_logged_in", None)
    return redirect(url_for("admin_login"))


@app.route("/admin/delete/<path:submission_date>", methods=["POST"])
@admin_required
def admin_delete(submission_date):
    if not CSV_PATH.exists():
        return redirect(url_for("admin_dashboard"))
    with open(CSV_PATH, newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    to_delete = [r for r in rows if r.get("Submission Date") == submission_date]
    remaining = [r for r in rows if r.get("Submission Date") != submission_date]
    for r in to_delete:
        if r.get("Word Doc"):
            (OUTPUT_DIR / r["Word Doc"]).unlink(missing_ok=True)
        if r.get("CV Filename"):
            (UPLOADS_DIR / r["CV Filename"]).unlink(missing_ok=True)
    with open(CSV_PATH, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=CSV_HEADERS, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(remaining)
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/download/doc/<path:filename>")
@admin_required
def admin_download_doc(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


@app.route("/admin/download/cv/<path:filename>")
@admin_required
def admin_download_cv(filename):
    return send_from_directory(UPLOADS_DIR, filename, as_attachment=True)


@app.route("/admin/jobs")
@admin_required
def admin_jobs():
    jobs = list(reversed(read_all_jobs()))
    candidates = read_all_candidates()
    assignments = read_assignments()
    assignments_by_job = {}
    for a in assignments:
        assignments_by_job.setdefault(a["Job ID"], []).append(a)
    clients = read_all_clients()
    return render_template("admin_jobs.html", jobs=jobs, candidates=candidates,
                           assignments_by_job=assignments_by_job, clients=clients)


@app.route("/admin/jobs/add", methods=["POST"])
@admin_required
def admin_add_job():
    f = request.form
    job = {
        "Job ID":           datetime.now().strftime("%Y%m%d%H%M%S"),
        "Date Added":       datetime.now().strftime("%Y-%m-%d %H:%M"),
        "Company Name":     f.get("company_name", "").strip(),
        "Job Title":        f.get("job_title", "").strip(),
        "Location":         _build_job_location(f.get("city","").strip(), f.get("country","").strip(), f.get("is_remote")=="yes"),
        "Salary Band":      f.get("salary_band", "").strip(),
        "Currency":         f.get("currency", "GBP").strip(),
        "Contact Person":   f.get("contact_person", "").strip(),
        "Assigned To":      f.get("assigned_to", "").strip(),
        "Status":           "Open",
        "Description":      f.get("description", "").strip(),
        "Show on Website":  "yes" if f.get("show_on_website") == "yes" else "no",
    }
    ensure_jobs_csv()
    with open(JOBS_PATH, "a", newline="", encoding="utf-8") as fh:
        csv.DictWriter(fh, fieldnames=JOBS_HEADERS, extrasaction="ignore").writerow(job)
    return redirect(url_for("admin_jobs"))


@app.route("/admin/jobs/status/<job_id>", methods=["POST"])
@admin_required
def admin_update_job_status(job_id):
    new_status = request.form.get("status", "Open")
    jobs = read_all_jobs()
    for j in jobs:
        if j["Job ID"] == job_id:
            j["Status"] = new_status
    with open(JOBS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=JOBS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(jobs)
    return redirect(url_for("admin_jobs"))


@app.route("/admin/jobs/<job_id>")
@admin_required
def admin_job_detail(job_id):
    import re as _re
    jobs = read_all_jobs()
    job = next((j for j in jobs if j["Job ID"] == job_id), None)
    if not job:
        return redirect(url_for("admin_jobs"))
    assignments = read_assignments()
    job_assignments = [a for a in assignments if a["Job ID"] == job_id]
    candidates = read_all_candidates()
    candidates_by_date = {c["Submission Date"]: c for c in candidates}
    clients = read_all_clients()
    client = next((c for c in clients if c["Client Name"] == job.get("Company Name", "")), None)

    # Currency symbol
    currency_code = job.get("Currency", "GBP") or "GBP"
    currency_symbol = CURRENCY_SYMBOLS.get(currency_code, currency_code + " ")

    # Calculate fee range from salary band
    fee_low = fee_high = None
    salary_band = job.get("Salary Band", "")
    if salary_band and client:
        try:
            fee_pct = float(client.get("Fee Percentage", 0))
        except (ValueError, TypeError):
            fee_pct = 0.0
        numbers = [float(_re.sub(r"[^\d.]", "", n)) for n in _re.findall(r"[\d,]+(?:\.\d+)?", salary_band) if _re.sub(r"[^\d.]", "", n)]
        if len(numbers) >= 2:
            fee_low  = numbers[0] * fee_pct / 100
            fee_high = numbers[1] * fee_pct / 100
        elif len(numbers) == 1:
            fee_low = fee_high = numbers[0] * fee_pct / 100

    return render_template(
        "admin_job_detail.html",
        job=job,
        assignments=job_assignments,
        candidates_by_date=candidates_by_date,
        stages=PIPELINE_STAGES,
        stage_colours=STAGE_COLOURS,
        client=client,
        fee_low=fee_low,
        fee_high=fee_high,
        currency_symbol=currency_symbol,
        currency_code=currency_code,
    )


@app.route("/admin/jobs/<job_id>/update", methods=["POST"])
@admin_required
def admin_update_job(job_id):
    jobs = read_all_jobs()
    for j in jobs:
        if j["Job ID"] == job_id:
            j["Status"]               = request.form.get("status", j.get("Status", "Open"))
            j["Currency"]             = request.form.get("currency", j.get("Currency", "GBP")).strip()
            j["Expected Start Date"]  = request.form.get("expected_start_date", "").strip()
            j["Expected Invoice Date"]= request.form.get("expected_invoice_date", "").strip()
    with open(JOBS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=JOBS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(jobs)
    return redirect(url_for("admin_job_detail", job_id=job_id))


@app.route("/admin/jobs/<job_id>/advert", methods=["POST"])
@admin_required
def admin_save_advert(job_id):
    jobs = read_all_jobs()
    for j in jobs:
        if j["Job ID"] == job_id:
            j["Advert Text"]    = request.form.get("advert_text", "").strip()
            j["Show on Website"]= "yes" if request.form.get("show_on_website") == "yes" else "no"
    with open(JOBS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=JOBS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(jobs)
    return redirect(url_for("admin_job_detail", job_id=job_id))


@app.route("/admin/jobs/visibility/<job_id>", methods=["POST"])
@admin_required
def admin_toggle_job_visibility(job_id):
    show = "yes" if request.form.get("show_on_website") == "yes" else "no"
    jobs = read_all_jobs()
    for j in jobs:
        if j["Job ID"] == job_id:
            j["Show on Website"] = show
    with open(JOBS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=JOBS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(jobs)
    return redirect(url_for("admin_jobs"))


@app.route("/admin/jobs/delete/<job_id>", methods=["POST"])
@admin_required
def admin_delete_job(job_id):
    jobs = read_all_jobs()
    with open(JOBS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=JOBS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(j for j in jobs if j["Job ID"] != job_id)
    assignments = read_assignments()
    with open(ASSIGNMENTS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=ASSIGNMENTS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(a for a in assignments if a["Job ID"] != job_id)
    return redirect(url_for("admin_jobs"))


@app.route("/admin/assign", methods=["POST"])
@admin_required
def admin_assign():
    job_id = request.form.get("job_id", "")
    candidate_date = request.form.get("candidate_date", "")
    candidate_name = request.form.get("candidate_name", "")
    ensure_assignments_csv()
    existing = read_assignments()
    already = any(a["Job ID"] == job_id and a["Candidate Date"] == candidate_date for a in existing)
    if not already and job_id and candidate_date:
        with open(ASSIGNMENTS_PATH, "a", newline="", encoding="utf-8") as fh:
            csv.DictWriter(fh, fieldnames=ASSIGNMENTS_HEADERS, extrasaction="ignore").writerow({
                "Job ID": job_id,
                "Candidate Date": candidate_date,
                "Candidate Name": candidate_name,
            })
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/unassign", methods=["POST"])
@admin_required
def admin_unassign():
    job_id = request.form.get("job_id", "")
    candidate_date = request.form.get("candidate_date", "")
    assignments = read_assignments()
    with open(ASSIGNMENTS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=ASSIGNMENTS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(a for a in assignments if not (a["Job ID"] == job_id and a["Candidate Date"] == candidate_date))
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/clients")
@admin_required
def admin_clients():
    clients = list(reversed(read_all_clients()))
    return render_template("admin_clients.html", clients=clients)


@app.route("/admin/clients/add", methods=["POST"])
@admin_required
def admin_add_client():
    f = request.form
    client_id = datetime.now().strftime("%Y%m%d%H%M%S")
    terms_filename = ""
    terms_file = request.files.get("terms_file")
    if terms_file and terms_file.filename:
        ext = Path(terms_file.filename).suffix.lower()
        terms_filename = f"terms_{client_id}{ext}"
        terms_file.save(TERMS_DIR / terms_filename)
    client = {
        "Client ID":      client_id,
        "Date Added":     datetime.now().strftime("%Y-%m-%d %H:%M"),
        "Client Name":    f.get("client_name", "").strip(),
        "Contact Name":   f.get("contact_name", "").strip(),
        "Fee Percentage": f.get("fee_percentage", "").strip(),
        "Terms Filename": terms_filename,
    }
    ensure_clients_csv()
    with open(CLIENTS_PATH, "a", newline="", encoding="utf-8") as fh:
        csv.DictWriter(fh, fieldnames=CLIENTS_HEADERS, extrasaction="ignore").writerow(client)
    return redirect(url_for("admin_clients"))


@app.route("/admin/clients/delete/<client_id>", methods=["POST"])
@admin_required
def admin_delete_client(client_id):
    clients = read_all_clients()
    to_delete = next((c for c in clients if c["Client ID"] == client_id), None)
    if to_delete and to_delete.get("Terms Filename"):
        (TERMS_DIR / to_delete["Terms Filename"]).unlink(missing_ok=True)
    with open(CLIENTS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=CLIENTS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(c for c in clients if c["Client ID"] != client_id)
    return redirect(url_for("admin_clients"))


@app.route("/admin/clients/terms/<path:filename>")
@admin_required
def admin_download_terms(filename):
    return send_from_directory(TERMS_DIR, filename, as_attachment=True)


@app.route("/admin/candidate/<path:candidate_date>")
@admin_required
def admin_candidate_detail(candidate_date):
    candidate_date = unquote(candidate_date)
    candidates = read_all_candidates()
    candidate = next((c for c in candidates if c.get("Submission Date") == candidate_date), None)
    if not candidate:
        return redirect(url_for("admin_dashboard"))
    tags = read_all_tags()
    candidate_tags = read_candidate_tags()
    c_tag_names = [ct["Tag Name"] for ct in candidate_tags if ct["Candidate Date"] == candidate_date]
    assignments = read_assignments()
    jobs = read_all_jobs()
    c_jobs = [j for a in assignments if a["Candidate Date"] == candidate_date
              for j in jobs if j["Job ID"] == a["Job ID"]]
    return render_template("admin_candidate_detail.html",
                           c=candidate, tags=tags, c_tag_names=c_tag_names,
                           c_jobs=c_jobs, jobs=jobs)


@app.route("/admin/candidate/<path:candidate_date>/update", methods=["POST"])
@admin_required
def admin_candidate_update(candidate_date):
    candidate_date = unquote(candidate_date)
    editable_fields = [
        "Full Name", "Email", "Phone", "Location", "LinkedIn URL",
        "Nearest Airport", "Current Job Title", "Current Company",
        "Currency", "Current Base Salary", "On Target Earnings",
        "Benefits", "Notice Period", "Open to Relocate",
        "Languages", "Education", "Data Consent Given",
    ]
    candidates = read_all_candidates()
    for row in candidates:
        if row.get("Submission Date") == candidate_date:
            for field in editable_fields:
                val = request.form.get(field)
                if val is not None:
                    row[field] = val.strip()
    with open(CSV_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=CSV_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(reversed(list(candidates)))
    return redirect(url_for("admin_candidate_detail", candidate_date=candidate_date))


@app.route("/admin/tags/add", methods=["POST"])
@admin_required
def admin_add_tag():
    tag_name = request.form.get("tag_name", "").strip()
    if tag_name:
        existing = read_all_tags()
        if not any(t["Tag Name"].lower() == tag_name.lower() for t in existing):
            colour = TAG_COLOURS[len(existing) % len(TAG_COLOURS)]
            ensure_tags_csv()
            with open(TAGS_PATH, "a", newline="", encoding="utf-8") as fh:
                csv.DictWriter(fh, fieldnames=TAGS_HEADERS, extrasaction="ignore").writerow({
                    "Tag ID": datetime.now().strftime("%Y%m%d%H%M%S"),
                    "Tag Name": tag_name,
                    "Colour": colour,
                })
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/tags/delete/<tag_id>", methods=["POST"])
@admin_required
def admin_delete_tag(tag_id):
    tags = read_all_tags()
    tag = next((t for t in tags if t["Tag ID"] == tag_id), None)
    with open(TAGS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=TAGS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(t for t in tags if t["Tag ID"] != tag_id)
    if tag:
        ct = read_candidate_tags()
        with open(CANDIDATE_TAGS_PATH, "w", newline="", encoding="utf-8") as fh:
            w = csv.DictWriter(fh, fieldnames=CANDIDATE_TAGS_HEADERS, extrasaction="ignore")
            w.writeheader()
            w.writerows(r for r in ct if r["Tag Name"] != tag["Tag Name"])
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/candidates/tag", methods=["POST"])
@admin_required
def admin_tag_candidate():
    candidate_date = request.form.get("candidate_date", "")
    tag_name = request.form.get("tag_name", "")
    redirect_to = request.form.get("redirect_to", "dashboard")
    if candidate_date and tag_name:
        ensure_candidate_tags_csv()
        existing = read_candidate_tags()
        already = any(r["Candidate Date"] == candidate_date and r["Tag Name"] == tag_name for r in existing)
        if not already:
            with open(CANDIDATE_TAGS_PATH, "a", newline="", encoding="utf-8") as fh:
                csv.DictWriter(fh, fieldnames=CANDIDATE_TAGS_HEADERS, extrasaction="ignore").writerow({
                    "Candidate Date": candidate_date,
                    "Tag Name": tag_name,
                })
    if redirect_to == "detail":
        return redirect(url_for("admin_candidate_detail", candidate_date=candidate_date))
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/candidates/untag", methods=["POST"])
@admin_required
def admin_untag_candidate():
    candidate_date = request.form.get("candidate_date", "")
    tag_name = request.form.get("tag_name", "")
    redirect_to = request.form.get("redirect_to", "dashboard")
    existing = read_candidate_tags()
    with open(CANDIDATE_TAGS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=CANDIDATE_TAGS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(r for r in existing if not (r["Candidate Date"] == candidate_date and r["Tag Name"] == tag_name))
    if redirect_to == "detail":
        return redirect(url_for("admin_candidate_detail", candidate_date=candidate_date))
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/jobs/<job_id>/pipeline")
@admin_required
def admin_pipeline(job_id):
    jobs = read_all_jobs()
    job = next((j for j in jobs if j["Job ID"] == job_id), None)
    if not job:
        return redirect(url_for("admin_jobs"))
    assignments = read_assignments()
    job_assignments = [a for a in assignments if a["Job ID"] == job_id]
    candidates = read_all_candidates()
    candidates_by_date = {c["Submission Date"]: c for c in candidates}
    return render_template("admin_pipeline.html", job=job,
                           assignments=job_assignments,
                           candidates_by_date=candidates_by_date,
                           stages=PIPELINE_STAGES,
                           stage_colours=STAGE_COLOURS)


@app.route("/admin/pipeline/stage", methods=["POST"])
@admin_required
def admin_pipeline_stage():
    job_id = request.form.get("job_id", "")
    candidate_date = request.form.get("candidate_date", "")
    stage = request.form.get("stage", "")
    assignments = read_assignments()
    for a in assignments:
        if a["Job ID"] == job_id and a["Candidate Date"] == candidate_date:
            a["Stage"] = stage
    with open(ASSIGNMENTS_PATH, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=ASSIGNMENTS_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(assignments)
    return redirect(url_for("admin_job_detail", job_id=job_id))


@app.route("/admin/pipeline/offer_letter", methods=["POST"])
@admin_required
def admin_pipeline_offer_letter():
    job_id = request.form.get("job_id", "")
    candidate_date = request.form.get("candidate_date", "")
    offer_file = request.files.get("offer_letter")
    if offer_file and offer_file.filename:
        ext = Path(offer_file.filename).suffix.lower()
        fname = f"offer_{job_id}_{safe_filename(candidate_date)}{ext}"
        offer_file.save(OFFER_LETTERS_DIR / fname)
        assignments = read_assignments()
        for a in assignments:
            if a["Job ID"] == job_id and a["Candidate Date"] == candidate_date:
                a["Offer Letter Filename"] = fname
        with open(ASSIGNMENTS_PATH, "w", newline="", encoding="utf-8") as fh:
            w = csv.DictWriter(fh, fieldnames=ASSIGNMENTS_HEADERS, extrasaction="ignore")
            w.writeheader()
            w.writerows(assignments)
    return redirect(url_for("admin_job_detail", job_id=job_id))


@app.route("/admin/pipeline/offer_letter/download/<path:filename>")
@admin_required
def admin_download_offer_letter(filename):
    return send_from_directory(OFFER_LETTERS_DIR, filename, as_attachment=True)


if __name__ == "__main__":
    ensure_csv()
    port = int(os.environ.get("PORT", 5050))
    app.run(debug=True, port=port)
